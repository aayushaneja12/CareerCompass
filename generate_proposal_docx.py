from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def add_field(paragraph, instruction_text: str):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction_text

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def add_page_number(footer_paragraph):
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.add_run("Page ")
    add_field(footer_paragraph, "PAGE")
    footer_paragraph.add_run(" of ")
    add_field(footer_paragraph, "NUMPAGES")


def add_heading_paragraph(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_numbered_points(doc: Document, points):
    for item in points:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(4)


def add_bulleted_points(doc: Document, points):
    for item in points:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)


def add_api_table(doc: Document):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Endpoint"
    hdr[1].text = "Method"
    hdr[2].text = "Purpose"
    hdr[3].text = "Auth"

    rows = [
        ("/health", "GET", "Service health check", "No"),
        ("/chat", "POST", "Conversational interaction via LangGraph", "Bearer JWT"),
        ("/profile", "GET", "Fetch current user career profile", "Bearer JWT"),
        ("/profile", "PUT", "Update current user career profile", "Bearer JWT"),
        ("/skill-gap", "POST", "Generate and persist skill gap report", "Bearer JWT"),
        ("/roadmap", "POST", "Generate and persist career roadmap", "Bearer JWT"),
        ("/roadmap/{roadmap_id}", "GET", "Fetch roadmap items", "Bearer JWT"),
        ("/resume-review", "POST", "Analyze resume and store review", "Bearer JWT"),
        ("/projects", "POST", "Generate and persist project ideas", "Bearer JWT"),
        ("/progress", "GET", "Fetch latest/specified weekly progress", "Bearer JWT"),
        ("/progress", "PUT", "Upsert weekly progress", "Bearer JWT"),
    ]

    for endpoint, method, purpose, auth in rows:
        cells = table.add_row().cells
        cells[0].text = endpoint
        cells[1].text = method
        cells[2].text = purpose
        cells[3].text = auth



def add_stack_table(doc: Document):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Layer"
    hdr[1].text = "Actual Technologies"
    hdr[2].text = "Rationale"

    rows = [
        (
            "Frontend",
            "React 18, TypeScript, Vite, React Router, Tailwind CSS, Radix UI",
            "Type-safe, modular SPA architecture with fast iteration and reusable UI primitives.",
        ),
        (
            "State and Data Access",
            "Custom hooks, Supabase JS client, shared authFetch API client",
            "Clear separation between UI concerns and async/network behavior.",
        ),
        (
            "Backend API",
            "FastAPI, Pydantic",
            "Typed request/response contracts and efficient async endpoint handling.",
        ),
        (
            "Agentic Orchestration",
            "LangGraph StateGraph, intent router, tool nodes",
            "Deterministic routing and composable tool execution for reliability.",
        ),
        (
            "LLM Integration",
            "LangChain ChatGroq (llama-3.3-70b-versatile)",
            "High-quality generative reasoning for structured career artifacts.",
        ),
        (
            "Database and Auth",
            "Supabase PostgreSQL, Supabase Auth, RLS-enabled career schema",
            "Managed auth + persistence with user-isolation policy model.",
        ),
        (
            "Dev Tooling",
            "ESLint, TypeScript, Uvicorn, SQL scripts",
            "Supports maintainability and local development workflows.",
        ),
    ]

    for layer, tech, rationale in rows:
        cells = table.add_row().cells
        cells[0].text = layer
        cells[1].text = tech
        cells[2].text = rationale



def build_document(output_path: Path):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Footer with page numbers
    section = doc.sections[0]
    footer_p = section.footer.paragraphs[0]
    add_page_number(footer_p)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("CareerCompass: Agentic AI Career Readiness Platform\n")
    r.bold = True
    r.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Comprehensive Technical and Academic Project Proposal")
    sr.font.size = Pt(14)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Author: Aayush\n").bold = True
    meta.add_run("Course: Capstone Project II (Placeholder where required)\n")
    meta.add_run("Program: Bachelor of Data Science\n")
    meta.add_run("Institution: SP Jain School of Global Management\n")
    meta.add_run(f"Date: {date.today().strftime('%d %B %Y')}")

    doc.add_page_break()

    # TOC page
    add_heading_paragraph(doc, "Table of Contents", level=1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    add_body_paragraph(doc, "Note: If the table appears unpopulated, right-click inside the table and choose 'Update Field' in Microsoft Word.")

    doc.add_page_break()

    # Sections
    add_heading_paragraph(doc, "1. Executive Summary", level=1)
    add_body_paragraph(
        doc,
        "CareerCompass is a full-stack AI-enabled career readiness platform that combines authenticated web workflows with an agentic conversational engine. "
        "The implemented system provides profile management, skill gap analysis, career roadmap generation, resume review, project recommendations, and weekly progress tracking, "
        "alongside PRP support capabilities such as events, FAQ retrieval, resource lookup, booking support, and conversation logging."
    )
    add_body_paragraph(
        doc,
        "The architecture uses a React and TypeScript frontend with Supabase-authenticated sessions, a FastAPI backend for typed service orchestration, "
        "LangGraph for intent-driven routing, and Supabase PostgreSQL for persistence. LLM functions are integrated through ChatGroq to generate structured career outputs. "
        "The result is a technically coherent platform that addresses both student guidance and operational scalability requirements."
    )

    add_heading_paragraph(doc, "2. Introduction", level=1)
    add_body_paragraph(
        doc,
        "Students frequently struggle to convert broad career aspirations into concrete plans. Existing tools are often fragmented, generic, and disconnected from each student's evolving profile and progress. "
        "CareerCompass addresses this by unifying conversational assistance and task-oriented modules within one authenticated product experience."
    )
    add_body_paragraph(
        doc,
        "The project demonstrates a practical academic-to-product transition: it implements real data persistence, authenticated workflows, and structured AI output generation rather than only static prototypes."
    )

    add_heading_paragraph(doc, "3. Problem Statement", level=1)
    add_body_paragraph(
        doc,
        "Career readiness support is often inconsistent across advising sessions, static content portals, and generic AI tools. Students lack continuity, role-specific diagnostics, and measurable progression pathways. "
        "Mentoring teams also spend substantial time on repetitive requests that can be partially automated."
    )
    add_body_paragraph(
        doc,
        "The core gap is the absence of an integrated, user-specific platform that connects profile data, intelligent analysis, structured planning, and weekly progress measurement in a secure and extensible architecture."
    )

    add_heading_paragraph(doc, "4. Objectives", level=1)
    add_heading_paragraph(doc, "4.1 General Objective", level=2)
    add_body_paragraph(doc, "To build a production-oriented academic system that provides personalized, scalable, and secure AI-assisted career readiness support.")

    add_heading_paragraph(doc, "4.2 Specific Objectives", level=2)
    add_numbered_points(doc, [
        "Implement authenticated career workflows for profile, analysis, and progress.",
        "Generate and persist role-aligned skill diagnostics and roadmaps.",
        "Provide resume quality and ATS-focused improvement feedback.",
        "Recommend portfolio projects mapped to target roles and user level.",
        "Enable conversational routing across PRP and career modules.",
    ])

    add_heading_paragraph(doc, "4.3 Technical Objectives", level=2)
    add_numbered_points(doc, [
        "Maintain type-safe API contracts between frontend and backend.",
        "Use deterministic graph routing before generative response composition.",
        "Persist generated artifacts for continuity and future analytics.",
        "Apply user-bound authorization for all protected operations.",
    ])

    add_heading_paragraph(doc, "5. Scope of the Project", level=1)
    add_heading_paragraph(doc, "5.1 In Scope", level=2)
    add_bulleted_points(doc, [
        "Protected React web application with multi-page feature navigation.",
        "FastAPI service layer with authenticated endpoints.",
        "LangGraph intent routing and tool execution pipeline.",
        "Supabase-backed persistence for conversations and career artifacts.",
        "LLM-backed generation for skill gap, roadmap, resume, and projects.",
    ])

    add_heading_paragraph(doc, "5.2 Out of Scope / Partial", level=2)
    add_bulleted_points(doc, [
        "Native mobile application.",
        "Enterprise-grade observability stack and SLO instrumentation.",
        "Fully automated PDF/DOCX extraction pipeline for resume text.",
        "Complete end-to-end automated test coverage for every module.",
    ])

    add_heading_paragraph(doc, "6. Proposed Solution", level=1)
    add_body_paragraph(
        doc,
        "CareerCompass delivers a dual interaction model. Users can complete structured workflows through dedicated pages (Profile, Skill Gap, Roadmap, Resume Review, Projects, Progress), "
        "or they can use conversational input through a chat interface that routes requests via intent classification and tool nodes."
    )
    add_body_paragraph(
        doc,
        "All protected requests include a Supabase access token. The backend verifies identity, executes business services, optionally invokes LLM generation, persists outcomes, and returns typed JSON responses to the frontend."
    )

    add_heading_paragraph(doc, "7. System Architecture", level=1)
    add_heading_paragraph(doc, "7.1 Frontend Architecture", level=2)
    add_body_paragraph(
        doc,
        "The frontend is built on React with TypeScript and route-level protection. Shared integration modules centralize API calls, and custom hooks encapsulate async state transitions, error handling, and loading behavior. "
        "Sidebar state is persisted in local storage, and chat view synchronizes with persisted messages."
    )

    add_heading_paragraph(doc, "7.2 Backend Architecture", level=2)
    add_body_paragraph(
        doc,
        "FastAPI exposes modular endpoints. Services are split by domain: profile service, career service, and resume service. "
        "A Supabase client is used for identity validation and database operations. Chat requests are delegated to a LangGraph workflow compiled once at startup."
    )

    add_heading_paragraph(doc, "7.3 Agentic Orchestration Layer", level=2)
    add_body_paragraph(
        doc,
        "The LangGraph flow starts at intent classification, then conditionally routes to a tool node (FAQ, booking, notes, events, knowledge search, profile, availability, skill gap, roadmap, resume, projects, mentor mode, progress, fallback). "
        "Tool outputs are passed to a response node for final user-facing generation, then logged to conversation history with title assignment."
    )

    add_heading_paragraph(doc, "7.4 Database and Auth Architecture", level=2)
    add_body_paragraph(
        doc,
        "Supabase Auth provides user sessions and access tokens. The backend verifies bearer tokens and binds operations to user identity. "
        "Data persistence spans baseline conversation entities and a dedicated career schema extension with row-level-security policies and indexed query paths."
    )

    add_heading_paragraph(doc, "7.5 End-to-End Flow", level=2)
    add_numbered_points(doc, [
        "User authenticates in frontend and receives a persisted session.",
        "Frontend sends protected API request with bearer token.",
        "Backend validates token and resolves user identity.",
        "Feature endpoint or graph route executes service logic.",
        "LLM generation occurs where needed and is parsed into structured output.",
        "Results are stored in Supabase and returned to UI for rendering.",
    ])

    add_heading_paragraph(doc, "8. Technology Stack", level=1)
    add_stack_table(doc)

    add_heading_paragraph(doc, "9. Methodology", level=1)
    add_body_paragraph(
        doc,
        "The implementation follows an iterative, modular, and API-driven methodology. The team established feature contracts using typed models, delivered backend and frontend modules in phases, "
        "and integrated AI generation through constrained JSON-prompt patterns to reduce hallucination risk and improve machine-parsable outputs."
    )
    add_body_paragraph(
        doc,
        "Prompt engineering strategy emphasizes role-context injection (for profile-aware analysis), strict output schemas, and fallback defaults when parsing fails. "
        "This combines generative flexibility with deterministic downstream processing."
    )

    add_heading_paragraph(doc, "10. Functional Modules", level=1)
    modules = [
        ("Authentication and Route Protection", "Secures all application pages except auth route using session checks and auth-state listeners."),
        ("Conversational Assistant", "Supports chat input, intent routing, tool execution, response synthesis, and conversation persistence."),
        ("Profile Management", "Provides CRUD-like profile lifecycle through get-or-create and update endpoints, with role/skill/goal fields."),
        ("Skill Gap Analysis", "Computes readiness score and prioritized missing skills for a target role and persists reports."),
        ("Career Roadmap", "Generates phased roadmap data and task items, then retrieves roadmap items for UI timeline rendering."),
        ("Resume Review", "Produces ATS score, strengths, weaknesses, keyword gaps, and improvement suggestions based on pasted resume text."),
        ("Project Suggestions", "Generates portfolio project ideas with difficulty, tech stack, relevance, and implementation steps."),
        ("Progress Tracking", "Stores weekly goals, accomplishments, challenges, learning minutes, and project completion indicators."),
        ("PRP Support Tools", "Includes FAQ retrieval, events lookup, KB snippet search, coach availability, bookings, notes, and follow-up advice."),
    ]
    for name, desc in modules:
        add_heading_paragraph(doc, name, level=2)
        add_body_paragraph(doc, desc)

    add_heading_paragraph(doc, "11. Agentic AI / Workflow Design", level=1)
    add_heading_paragraph(doc, "11.1 Intent Handling and Routing", level=2)
    add_body_paragraph(
        doc,
        "The classifier uses regex pattern sets to map user text to intents. A conditional router selects the corresponding node. If no match is found, fallback guidance is returned."
    )
    add_heading_paragraph(doc, "11.2 State Model", level=2)
    add_body_paragraph(
        doc,
        "GraphState carries messages, intent, user and conversation identifiers, and career-specific fields such as selected role, profile cache, skill-gap results, roadmap data, resume review results, and project suggestions."
    )
    add_heading_paragraph(doc, "11.3 Response Synthesis", level=2)
    add_body_paragraph(
        doc,
        "After tool execution, a response node calls ChatGroq with a strongly defined mentoring system prompt plus tool output context. This preserves factual grounding while improving conversational quality."
    )

    add_heading_paragraph(doc, "12. API Design", level=1)
    add_body_paragraph(doc, "The implemented API surface is summarized below.")
    add_api_table(doc)
    add_body_paragraph(
        doc,
        "Frontend integration uses a shared authFetch utility that retrieves or refreshes the Supabase session token, adds authorization headers, standardizes JSON handling, and surfaces user-friendly errors."
    )

    add_heading_paragraph(doc, "13. Database Design", level=1)
    add_heading_paragraph(doc, "13.1 Core Entities", level=2)
    add_bulleted_points(doc, [
        "profiles",
        "conversations",
        "messages",
    ])
    add_heading_paragraph(doc, "13.2 Career Extension Entities", level=2)
    add_bulleted_points(doc, [
        "user_profiles",
        "roadmaps",
        "roadmap_items",
        "skill_gap_reports",
        "resume_reviews",
        "saved_projects",
        "progress_metrics",
    ])
    add_heading_paragraph(doc, "13.3 Security and Integrity", level=2)
    add_body_paragraph(
        doc,
        "The career schema defines row-level-security policies that constrain access to each authenticated user's records. "
        "Indexes support common query paths such as user-specific roadmap retrieval and weekly progress lookup."
    )

    add_heading_paragraph(doc, "14. Technical Specifications", level=1)
    add_numbered_points(doc, [
        "Architecture style: modular full-stack web system with graph-orchestrated conversational backend.",
        "State handling: frontend local/hook state + backend GraphState + persisted Supabase records.",
        "Authentication: Supabase JWT bearer token verification at API boundary.",
        "Persistence: direct Supabase table operations for both transactional and generated artifacts.",
        "Error handling: endpoint exceptions, typed frontend catch paths, JSON parse fallbacks in AI services.",
        "Scalability path: stateless API and managed DB/Auth services support horizontal backend scaling.",
    ])

    add_heading_paragraph(doc, "15. Implementation Details", level=1)
    add_body_paragraph(
        doc,
        "Key implementation includes protected routes in the SPA, dedicated feature pages and hooks, backend endpoint expansion for career workflows, graph routing updates, and service-layer LLM integrations."
    )
    add_body_paragraph(
        doc,
        "Notable engineering choices include normalizing current_role/current_position key differences across model and database boundaries, "
        "persisting generated roadmap tasks as separate relational items, and centralizing auth header creation for all feature APIs."
    )

    add_heading_paragraph(doc, "16. Security and Privacy", level=1)
    add_bulleted_points(doc, [
        "Authenticated-only access to protected pages and APIs.",
        "Bearer token validation via Supabase auth get_user calls.",
        "User-scoped persistence model with RLS policies in career tables.",
        "No exposure of backend service key in frontend runtime.",
    ])
    add_body_paragraph(
        doc,
        "Operational risk remains around schema consistency and privilege handling if service-role write patterns are expanded. "
        "Mitigation requires strict endpoint-level user binding and periodic schema-contract verification."
    )

    add_heading_paragraph(doc, "17. Challenges and Solutions", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light List Accent 1"
    table.rows[0].cells[0].text = "Challenge"
    table.rows[0].cells[1].text = "Implemented Solution"
    challenges = [
        ("Combining chat and structured module flows", "Introduced parallel interaction model with shared auth and persistence semantics."),
        ("LLM output unpredictability", "Forced JSON output prompts with parser fallback defaults."),
        ("Field naming drift between layers", "Added mapping and normalization functions in service/model adapters."),
        ("Conversation continuity", "Persisted conversation/message records and auto-generated titles."),
        ("Incremental frontend expansion", "Built feature-specific hooks and API modules to avoid monolithic component logic."),
    ]
    for c, s in challenges:
        row = table.add_row().cells
        row[0].text = c
        row[1].text = s

    add_heading_paragraph(doc, "18. Testing and Validation", level=1)
    add_body_paragraph(
        doc,
        "Validation should cover endpoint authorization, feature correctness, UI interaction flows, data persistence verification, and LLM fallback reliability. "
        "Implemented code patterns already support deterministic verification of API contracts and state transitions, while complete automated test expansion remains a recommended next step."
    )
    add_numbered_points(doc, [
        "Functional scenario tests for each feature module.",
        "Auth failure and token expiry tests.",
        "Database side-effect verification after API calls.",
        "Chat intent routing and fallback behavior validation.",
        "Frontend flow testing for route guards and error states.",
    ])

    add_heading_paragraph(doc, "19. Results and Expected Outcomes", level=1)
    add_body_paragraph(
        doc,
        "The implemented system delivers a cohesive student-facing product with real feature depth beyond prototype-level claims. "
        "Expected outcomes include improved user clarity on career trajectories, reduced mentor overhead on repetitive informational tasks, and stronger evidence-driven planning through persisted progress records."
    )

    add_heading_paragraph(doc, "20. Future Enhancements", level=1)
    add_numbered_points(doc, [
        "End-to-end automated test suites across frontend, backend, and graph routing.",
        "Unified schema/type generation pipeline to remove field drift.",
        "Document parsing for PDF and DOCX resume ingestion.",
        "Analytics dashboards for readiness trends and program-level insights.",
        "Improved memory and personalization layers for long-horizon mentoring.",
        "Production observability, monitoring, and CI/CD hardening.",
        "Mobile-first and multi-device UX optimization.",
    ])

    add_heading_paragraph(doc, "21. Conclusion", level=1)
    add_body_paragraph(
        doc,
        "CareerCompass demonstrates a mature capstone implementation that integrates modern web engineering, secure authentication, deterministic workflow orchestration, and practical AI augmentation. "
        "Its architecture is both academically rigorous and product-credible, making it a strong foundation for continued institutional deployment and research-driven enhancement."
    )

    add_heading_paragraph(doc, "22. Diagrams Section (Diagram-Ready Content)", level=1)

    diagrams = [
        (
            "22.1 System Architecture Diagram",
            [
                "Components: User, React Frontend, Supabase Auth, FastAPI Backend, LangGraph, Tool Nodes, Career Services, ChatGroq LLM, Supabase PostgreSQL.",
                "Relationships: Frontend authenticates via Supabase, sends bearer-authenticated requests to backend, backend routes through LangGraph/tool nodes/services, persists to DB, and returns response.",
                "Explanation: This illustrates layered interaction from UI to agentic processing and persistence.",
                "Mermaid:",
                "flowchart LR",
                "U[User] --> FE[React Frontend]",
                "FE --> SA[Supabase Auth]",
                "FE --> API[FastAPI]",
                "API --> LG[LangGraph]",
                "LG --> TN[Tool Nodes]",
                "TN --> CS[Career Services]",
                "CS --> LLM[ChatGroq]",
                "CS --> DB[(Supabase PostgreSQL)]",
                "LG --> DB",
                "API --> FE",
            ],
        ),
        (
            "22.2 Data Flow Diagram",
            [
                "Components: Input Forms/Chat, authFetch client, API endpoints, service layer, DB tables, response renderers.",
                "Relationships: User input -> authenticated request -> service logic -> DB write/read -> structured response -> UI state update.",
                "Explanation: Shows end-to-end data movement for both chat and feature modules.",
            ],
        ),
        (
            "22.3 Agent Workflow / Graph Diagram",
            [
                "Components: classify_intent, router, tool nodes, response_node, rewrite_node, conversation_log_node, conversation_title_node.",
                "Relationships: classify -> route -> tool -> response -> rewrite -> log -> title -> end.",
                "Explanation: Demonstrates deterministic orchestration plus LLM-based response synthesis.",
                "Mermaid:",
                "flowchart TD",
                "A[classify_intent] --> B{router}",
                "B --> C1[faq/booking/notes/events/kb/profile/availability]",
                "B --> C2[skill_gap/roadmap/resume/projects/mentor/progress]",
                "B --> X[fallback]",
                "C1 --> R[response_node]",
                "C2 --> R",
                "X --> R",
                "R --> W[rewrite_node] --> L[conversation_log_node] --> T[conversation_title_node] --> E[END]",
            ],
        ),
        (
            "22.4 Frontend-Backend Interaction Diagram",
            [
                "Components: Browser, Supabase Auth, API client, FastAPI, Supabase DB.",
                "Relationships: Browser obtains token, api-client sends bearer request, backend validates and processes, DB persists, response returns.",
                "Explanation: Captures authenticated contract between client and server.",
            ],
        ),
        (
            "22.5 User Flow Diagram",
            [
                "Components: Sign-in, Dashboard, Feature Selection, Analysis Results, Progress Updates.",
                "Relationships: Auth -> chat or module -> generated insights -> persisted data -> iterative improvement loop.",
                "Explanation: Illustrates user journey from access to continuous career development.",
            ],
        ),
    ]

    for title, lines in diagrams:
        add_heading_paragraph(doc, title, level=2)
        for line in lines:
            add_body_paragraph(doc, line)

    add_heading_paragraph(doc, "Appendix A: Verified Artifacts from Code Inspection", level=1)
    add_bulleted_points(doc, [
        "Frontend routes and pages implemented for profile, skill-gap, roadmap, resume, projects, and progress.",
        "Backend app includes chat and career endpoint set with bearer-auth checks.",
        "LangGraph includes registered intent classifier, tool nodes, response, logging, and title nodes.",
        "Career schema SQL defines seven extension tables with RLS and index setup.",
        "Frontend integration modules include API client, chat, profile, career, resume, and progress adapters.",
    ])

    doc.save(str(output_path))


if __name__ == "__main__":
    output = Path("CareerCompass_Project_Proposal_Aayush.docx")
    build_document(output)
    print(f"Generated: {output.resolve()}")
